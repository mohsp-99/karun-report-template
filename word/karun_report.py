"""
Karun branded report engine.

Renders a structured report (a Python dict / JSON file) into a branded Word
.docx that mirrors the Karun LaTeX template (karun.sty + titlepage.tex):

  - Full-bleed cover page: background swoosh, top-right logo, title/subtitle,
    metadata table with strike-through Access Level / Confidentiality toggles,
    footer logo + disclaimer + copyright.
  - Running header (date | summary title) and footer (Page X of Y | logo),
    each with a Karun-blue rule.
  - Karun-blue numbered headings (1, 1.1) and an auto Table of Contents.
  - Body paragraphs with lightweight inline markup, bullet lists, figures
    with captions.

The chrome is built ONCE here, correctly, and is data-driven. An AI assistant
only has to produce the content JSON (see SKILL.md); it never touches Word
internals.

Usage:
    python karun_report.py <content.json> <output.docx>
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

ASSETS = Path(__file__).resolve().parent / "assets"

# ---------------------------------------------------------------------------
# Palette (from karun.sty \definecolor blocks)
# ---------------------------------------------------------------------------
KARUN_BLUE = RGBColor(0x00, 0x32, 0x64)   # RGB(0,50,100)
GRAY_SUB = RGBColor(0x80, 0x80, 0x80)     # subtitle gray
BODY_BLACK = RGBColor(0x00, 0x00, 0x00)
STRIKE_GRAY = RGBColor(0x78, 0x78, 0x78)  # access-strike RGB(120,120,120)

ACCESS_COLORS = {                          # active option -> color
    1: RGBColor(0xB4, 0xB4, 0xB4),         # Internal  -> access-gray
    2: RGBColor(0x00, 0x66, 0x99),         # External  -> access-blue
    3: RGBColor(0x00, 0x80, 0x00),         # Public    -> access-green
}
ACCESS_LABELS = {"en": ["Internal", "External", "Public"]}

CONF_COLORS = {
    1: RGBColor(0x00, 0x80, 0x00),         # Normal       -> green
    2: RGBColor(0xFF, 0xA5, 0x00),         # Confidential -> orange
    3: RGBColor(0xFF, 0x00, 0x00),         # Top Secret   -> red
}
CONF_LABELS = {"en": ["Normal", "Confidential", "Top Secret"]}

FONT = "Dubai"          # main template font (XeLaTeX \setmainfont{Dubai})
A4_W = Cm(21.0)
A4_H = Cm(29.7)

BLUE_HEX = "003264"


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------
def _set(el, **attrs):
    for k, v in attrs.items():
        el.set(qn(k if ":" in k else f"w:{k}"), str(v))
    return el


def _sub(parent, tag, **attrs):
    el = OxmlElement(tag)
    _set(el, **attrs)
    parent.append(el)
    return el


def _setp(el, **attrs):
    """Set unqualified (no-namespace) attributes, as required by DrawingML."""
    for k, v in attrs.items():
        el.set(k, str(v))
    return el


def _subp(parent, tag, **attrs):
    el = OxmlElement(tag)
    _setp(el, **attrs)
    parent.append(el)
    return el


def add_field(paragraph, instr, *, size=9, font=FONT, color=None, bold=False):
    """Insert a Word field (e.g. PAGE, SECTIONPAGES, TOC) into a paragraph."""
    run = paragraph.add_run()
    _apply_run_format(run, size=size, font=font, color=color, bold=bold)
    r = run._r
    fld_begin = _sub(r, "w:fldChar", fldCharType="begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    r.append(instr_el)
    _sub(r, "w:fldChar", fldCharType="separate")
    _sub(r, "w:fldChar", fldCharType="end")
    return run


def _apply_run_format(run, *, size=None, font=FONT, color=None, bold=False,
                      italic=False, code=False, strike=False):
    if code:
        run.font.name = "Consolas"
    elif font:
        run.font.name = font
        rPr = run._element.get_or_add_rPr()
        rfonts = rPr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = _sub(rPr, "w:rFonts")
        _set(rfonts, ascii=font, hAnsi=font, cs=font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.strike = strike
    if color is not None:
        run.font.color.rgb = color
    return run


def paragraph_border(paragraph, edge, *, color=BLUE_HEX, size=8, space=4):
    """Add a single coloured rule on one edge of a paragraph (header/footer)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    _sub(pBdr, f"w:{edge}", val="single", sz=size, space=space, color=color)


def add_floating_image(paragraph, image_path, *, width_cm, height_cm=None,
                       pos_x_cm=0.0, pos_y_cm=0.0, rel="page", behind=False):
    """Anchor an image at an absolute position relative to the page.

    Used for the cover background, cover logos and footer logo.
    """
    run = paragraph.add_run()
    # Let python-docx ingest the image & build the inline drawing first, then
    # convert that inline drawing into a floating anchor.
    from docx.shared import Cm as _Cm
    if height_cm is not None:
        run.add_picture(str(image_path), width=_Cm(width_cm), height=_Cm(height_cm))
    else:
        run.add_picture(str(image_path), width=_Cm(width_cm))
    inline = run._r.find(qn("w:drawing")).find(qn("wp:inline"))
    extent = inline.find(qn("wp:extent"))
    cx, cy = extent.get("cx"), extent.get("cy")

    # Pull the graphic payload out of the inline element.
    graphic = inline.find(qn("a:graphic"))
    doc_pr = inline.find(qn("wp:docPr"))

    anchor = OxmlElement("wp:anchor")
    _setp(anchor, distT="0", distB="0", distL="0", distR="0",
          simplePos="0", relativeHeight="0" if behind else "5",
          behindDoc="1" if behind else "0", locked="0",
          layoutInCell="1", allowOverlap="1")
    _subp(anchor, "wp:simplePos", x="0", y="0")
    posH = _subp(anchor, "wp:positionH", relativeFrom=rel)
    _subp(posH, "wp:posOffset").text = str(int(Cm(pos_x_cm)))
    posV = _subp(anchor, "wp:positionV", relativeFrom=rel)
    _subp(posV, "wp:posOffset").text = str(int(Cm(pos_y_cm)))
    ext = _subp(anchor, "wp:extent", cx=cx, cy=cy)
    _subp(anchor, "wp:effectExtent", l="0", t="0", r="0", b="0")
    _subp(anchor, "wp:wrapNone")
    anchor.append(doc_pr)
    cnvGr = _subp(anchor, "wp:cNvGraphicFramePr")
    _subp(cnvGr, "a:graphicFrameLocks", noChangeAspect="1")
    anchor.append(graphic)

    drawing = run._r.find(qn("w:drawing"))
    drawing.remove(inline)
    drawing.append(anchor)
    return run


# ---------------------------------------------------------------------------
# Inline markup: **bold**, *italic*, `code`
# ---------------------------------------------------------------------------
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", re.DOTALL)


def add_inline(paragraph, text, *, size=11, color=BODY_BLACK, font=FONT):
    """Render a string with lightweight markdown-style inline formatting."""
    if not text:
        return
    for chunk in _INLINE_RE.split(text):
        if not chunk:
            continue
        bold = italic = code = False
        body = chunk
        if chunk.startswith("**") and chunk.endswith("**"):
            bold, body = True, chunk[2:-2]
        elif chunk.startswith("*") and chunk.endswith("*"):
            italic, body = True, chunk[1:-1]
        elif chunk.startswith("`") and chunk.endswith("`"):
            code, body = True, chunk[1:-1]
        run = paragraph.add_run(body)
        _apply_run_format(run, size=size, font=font, color=color,
                          bold=bold, italic=italic, code=code)


# ---------------------------------------------------------------------------
# Document styles
# ---------------------------------------------------------------------------
def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_BLACK
    rPr = normal.element.get_or_add_rPr()
    rfonts = rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = _sub(rPr, "w:rFonts")
    _set(rfonts, ascii=FONT, hAnsi=FONT, cs=FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, sz in (("Heading 1", 20), ("Heading 2", 14), ("Heading 3", 12)):
        st = styles[name]
        st.font.name = FONT
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = KARUN_BLUE
        st.paragraph_format.space_before = Pt(14)
        st.paragraph_format.space_after = Pt(10)
        st.paragraph_format.keep_with_next = True
        # Stop Word from auto-numbering; we control numbers ourselves.


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------
def build_cover(doc, meta, lang="en"):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = A4_W, A4_H
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(2.0))
    sec.different_first_page_header_footer = True
    logo = ASSETS / ("karun-en.png" if lang == "en" else "karun-fa.png")

    # Background swoosh: full-page floating image behind everything, placed in
    # the first-page header so it does not push body content.
    hdr = sec.first_page_header
    bg_p = hdr.paragraphs[0]
    add_floating_image(bg_p, ASSETS / "bg.jpeg", width_cm=21.0, height_cm=29.7,
                       pos_x_cm=0.0, pos_y_cm=0.0, rel="page", behind=True)
    # Top-right logo (4 cm), 1 cm inset.
    add_floating_image(bg_p, logo, width_cm=4.0,
                       pos_x_cm=21.0 - 1.0 - 4.0, pos_y_cm=1.0, rel="page")
    # Bottom-left logo (2.5 cm) on the cover.
    add_floating_image(bg_p, logo, width_cm=2.5,
                       pos_x_cm=1.0, pos_y_cm=29.7 - 2.0, rel="page")

    body = doc.add_paragraph()
    body.paragraph_format.space_before = Pt(0)
    # Push title down to ~5 cm.
    spacer = body
    spacer.paragraph_format.space_before = Cm(4.5)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_after = Pt(6)
    r = title_p.add_run(meta["title"])
    _apply_run_format(r, size=30, bold=True, color=BODY_BLACK)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Cm(1.3)
    r = sub_p.add_run(meta.get("subtitle", ""))
    _apply_run_format(r, size=15, color=GRAY_SUB)

    build_meta_table(doc, meta, lang)

    # Cover footer: small logo bottom-left, disclaimer, copyright.
    foot = sec.first_page_footer
    fp = foot.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.left_indent = Cm(3.2)  # clear the bottom-left logo
    add_inline(
        fp,
        "This document is prepared by the Karun team for the intended "
        "recipients mentioned herein. Any reuse or redistribution to third "
        "parties, in any form, is prohibited without prior written consent.",
        size=7, color=GRAY_SUB,
    )
    cr = foot.add_paragraph()
    cr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = cr.add_run(f"© {meta.get('year', 2025)} Karun, Iran")
    _apply_run_format(r, size=9, color=BODY_BLACK)


def build_meta_table(doc, meta, lang="en"):
    rows = [
        ("Document Title:", [("b", meta.get("summary_title", ""))]),
        (meta.get("client_label", "Client:"), [("n", meta.get("client", ""))]),
        ("Prepared by:", [("n", meta.get("prepared_by", ""))]),
        ("Access Level:", _toggle_runs(meta.get("access_level", 1),
                                       ACCESS_LABELS[lang], ACCESS_COLORS)),
        ("Confidentiality:", _toggle_runs(meta.get("confidentiality", 1),
                                          CONF_LABELS[lang], CONF_COLORS)),
        ("Document ID:", [("n", meta.get("doc_id", ""))]),
        ("Date:", [("n", meta.get("date", ""))]),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.allow_autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(11.8)
    _table_top_bottom_borders(table)
    for i, (label, value_runs) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.width, c1.width = Cm(4.2), Cm(11.8)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after = Pt(4)
        r = p0.add_run(label)
        _apply_run_format(r, size=11, bold=True, color=BODY_BLACK)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after = Pt(4)
        for kind, *rest in value_runs:
            if kind == "sep":
                run = p1.add_run(" / ")
                _apply_run_format(run, size=11, color=BODY_BLACK)
            else:
                text = rest[0]
                color = rest[1] if len(rest) > 1 else BODY_BLACK
                run = p1.add_run(text)
                _apply_run_format(
                    run, size=11, color=color,
                    bold=(kind in ("b", "active")),
                    strike=(kind == "strike"),
                )


def _toggle_runs(active, labels, color_map):
    """Build the strike-through option row: active is bold/coloured, the rest
    are struck through in gray, joined by ' / '."""
    out = []
    for idx, label in enumerate(labels, start=1):
        if idx > 1:
            out.append(("sep",))
        if idx == active:
            out.append(("active", label, color_map[active]))
        else:
            out.append(("strike", label, STRIKE_GRAY))
    return out


def _table_top_bottom_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        sz = 18 if edge in ("top", "bottom") else 6
        _sub(borders, f"w:{edge}", val="single", sz=sz, space="0", color="000000")
    for edge in ("left", "right"):
        _sub(borders, f"w:{edge}", val="none", sz="0", space="0", color="auto")
    tblPr.append(borders)


# ---------------------------------------------------------------------------
# Running header / footer (body section)
# ---------------------------------------------------------------------------
def setup_running_header_footer(section, meta, logo):
    section.different_first_page_header_footer = False

    hdr = section.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tabs = hp.paragraph_format.tab_stops
    usable = section.page_width - section.left_margin - section.right_margin
    tabs.add_tab_stop(Emu(int(usable / 2)), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Emu(int(usable)), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run(meta.get("date", ""))
    _apply_run_format(r, size=9, color=BODY_BLACK)
    r = hp.add_run("\t" + meta.get("summary_title", ""))
    _apply_run_format(r, size=9, color=BODY_BLACK)
    paragraph_border(hp, "bottom", color=BLUE_HEX, size=8, space=6)

    ftr = section.footer
    ftr.is_linked_to_previous = False
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tabs = fp.paragraph_format.tab_stops
    tabs.add_tab_stop(Emu(int(usable)), WD_TAB_ALIGNMENT.RIGHT)
    r = fp.add_run("Page ")
    _apply_run_format(r, size=9, color=BODY_BLACK)
    add_field(fp, "PAGE", size=9)
    r = fp.add_run(" of ")
    _apply_run_format(r, size=9, color=BODY_BLACK)
    # SECTIONPAGES (this section only), not NUMPAGES (whole doc): the body is its
    # own section with page numbering restarted at 1, so this excludes the cover
    # and the last page reads "Page n of n" instead of "Page n-1 of n".
    add_field(fp, "SECTIONPAGES", size=9)
    fp.add_run("\t")
    logo_run = fp.add_run()
    logo_run.add_picture(str(logo), width=Cm(2.2))
    paragraph_border(fp, "top", color=BLUE_HEX, size=8, space=6)


# ---------------------------------------------------------------------------
# Body rendering
# ---------------------------------------------------------------------------
class Numberer:
    def __init__(self):
        self.s = 0
        self.ss = 0

    def section(self):
        self.s += 1
        self.ss = 0
        return str(self.s)

    def subsection(self):
        self.ss += 1
        return f"{self.s}.{self.ss}"


def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run("Contents")
    _apply_run_format(r, size=18, bold=True, color=KARUN_BLUE)
    p.paragraph_format.space_after = Pt(12)
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u', size=11)


def add_figure(doc, item, base_dir):
    img = _resolve_image(item["image"], base_dir)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    usable = A4_W - Cm(3.1) - Cm(2.5)
    run.add_picture(str(img), width=usable)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(12)
    label = item.get("caption_label", "Figure")
    r = cap.add_run(f"{label}: ")
    _apply_run_format(r, size=9, bold=True, italic=True, color=KARUN_BLUE)
    add_inline(cap, item.get("caption", ""), size=9, color=BODY_BLACK)
    for r in cap.runs[1:]:
        r.font.italic = True
        r.font.size = Pt(9)


def _resolve_image(name, base_dir):
    p = Path(name)
    if p.is_absolute() and p.exists():
        return p
    for cand in (base_dir / name, base_dir / "images" / name, ASSETS / name):
        if cand.exists():
            return cand
    return base_dir / name


def render_body(doc, body, base_dir):
    num = Numberer()
    for item in body:
        # Skip comment/placeholder items (no "type", or keys like "_comment").
        if "type" not in item:
            continue
        t = item["type"]
        if t == "heading":
            level = item.get("level", 1)
            text = item["text"]
            numbered = item.get("numbered", True)
            if level == 1:
                prefix = (num.section() + "  ") if numbered else ""
                p = doc.add_paragraph(style="Heading 1")
            else:
                prefix = (num.subsection() + "  ") if numbered else ""
                p = doc.add_paragraph(style="Heading 2")
            run = p.add_run(prefix + text)
        elif t == "paragraph":
            p = doc.add_paragraph()
            add_inline(p, item.get("text", ""), size=11)
        elif t == "list":
            for it in item["items"]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_inline(p, it, size=11)
        elif t == "figure":
            add_figure(doc, item, base_dir)
        elif t == "pagebreak":
            doc.add_page_break()
        else:
            raise ValueError(f"Unknown body item type: {t!r}")


# ---------------------------------------------------------------------------
# Top-level build
# ---------------------------------------------------------------------------
def build(report, out_path, base_dir):
    lang = report.get("lang", "en")
    meta = report["meta"]
    doc = Document()
    setup_styles(doc)
    build_cover(doc, meta, lang)

    # New section for the running body (its own header/footer, restart page #).
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    body_sec.page_width, body_sec.page_height = A4_W, A4_H
    body_sec.top_margin = Cm(2.4)
    body_sec.bottom_margin = Cm(3.0)
    body_sec.left_margin = Cm(3.1)
    body_sec.right_margin = Cm(2.5)
    body_sec.header_distance = Cm(0.8)
    body_sec.footer_distance = Cm(1.2)
    # Restart page numbering at 1 for this section.
    sectPr = body_sec._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)

    logo = ASSETS / ("karun-en.png" if lang == "en" else "karun-fa.png")
    setup_running_header_footer(body_sec, meta, logo)

    add_toc(doc)
    doc.add_page_break()
    render_body(doc, report["body"], base_dir)

    doc.save(str(out_path))
    return out_path


def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    content_path = Path(sys.argv[1]).resolve()
    out_path = Path(sys.argv[2]).resolve()
    report = json.loads(content_path.read_text(encoding="utf-8"))
    build(report, out_path, content_path.parent)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
